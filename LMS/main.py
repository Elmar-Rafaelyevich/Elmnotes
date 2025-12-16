import telebot
import json
from telebot import types
from datetime import datetime
from io import BytesIO
from docx import Document
import speech_recognition as sr
from pydub import AudioSegment
from pydub.utils import which
import os

from config import TOKEN
from data import help_text, is_not_valid_group

AudioSegment.converter = which("ffmpeg")

bot = telebot.TeleBot(TOKEN)
GROUPS_FILE = os.path.join(os.path.dirname(__file__), '..', 'group', 'group.json')

# Состояние пользователя: group, subject, action
user_state = {}

# ==================== START ====================
@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(message.chat.id, text=help_text)
    bot.register_next_step_handler(message, get_group)

# ==================== ГРУППА ====================
def get_group(message):
    group_name = message.text.strip()

    with open(GROUPS_FILE, "r", encoding="utf-8") as f:
        schedules = json.load(f)

    if group_name not in schedules:
        bot.send_message(message.chat.id, text=is_not_valid_group)
        bot.register_next_step_handler(message, get_group)
        return

    user_state[message.chat.id] = {
        "group": group_name,
        "subject": None,
        "action": None
    }

    send_today_schedule(message, group_name)

# ==================== РАСПИСАНИЕ ====================
def send_today_schedule(message, group_name):
    with open(GROUPS_FILE, "r", encoding="utf-8") as f:
        schedules = json.load(f)

    days_map = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
    today_index = datetime.today().weekday()
    today = days_map[today_index]

    subjects_today = schedules[group_name].get(today, [])
    if not subjects_today:
        bot.send_message(message.chat.id, f"На {today} у группы {group_name} нет занятий 😴")
        return

    markup = types.InlineKeyboardMarkup()
    for subject in subjects_today:
        markup.add(types.InlineKeyboardButton(text=f"📚 {subject}", callback_data=f"subject|{subject}"))

    bot.send_message(message.chat.id, f"🗓 Сегодня {today}.\nВыбери предмет 📚", reply_markup=markup)

# ==================== ВЫБОР ПРЕДМЕТА ====================
@bot.callback_query_handler(func=lambda call: call.data.startswith("subject|"))
def callback_subject(call):
    subject = call.data.split("|")[1]
    user_state[call.message.chat.id]["subject"] = subject

    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Написать лекцию ✏️", callback_data="write_lecture"))
    markup.add(types.InlineKeyboardButton("Загрузить аудио 🎤", callback_data="upload_audio"))

    bot.send_message(
        call.message.chat.id,
        f"📘 Вы выбрали предмет: {subject}\nЧто хотите сделать❓",
        reply_markup=markup
    )

# ==================== НАПИСАТЬ ЛЕКЦИЮ ====================
@bot.callback_query_handler(func=lambda call: call.data == "write_lecture")
def write_lecture_handler(call):
    chat_id = call.message.chat.id
    user_state[chat_id]["action"] = "write_lecture"
    bot.send_message(chat_id, "Отправьте текст лекции ✏️")
    bot.register_next_step_handler_by_chat_id(chat_id, save_lecture)

def save_lecture(message):
    chat_id = message.chat.id
    subject = user_state[chat_id]["subject"]
    lecture_text = message.text

    # TXT
    txt_file = BytesIO()
    txt_file.write(lecture_text.encode("utf-8"))
    txt_file.seek(0)

    # DOCX
    doc = Document()
    doc.add_paragraph(lecture_text)
    doc_file = BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)

    bot.send_document(chat_id, ("{}.txt".format(subject), txt_file))
    bot.send_document(chat_id, ("{}.docx".format(subject), doc_file))
    bot.send_message(chat_id, f"📝 Лекции по предмету {subject} отправлены ✅")

    # После отправки снова показываем расписание
    send_today_schedule(message, user_state[chat_id]["group"])

# ==================== ЗАГРУЗИТЬ АУДИО ====================
@bot.callback_query_handler(func=lambda call: call.data == "upload_audio")
def upload_audio_handler(call):
    chat_id = call.message.chat.id
    user_state[chat_id]["action"] = "upload_audio"
    bot.send_message(chat_id, "🔊 Отправьте аудио файл (mp3, wav или голосовое сообщение)")
    bot.register_next_step_handler_by_chat_id(chat_id, process_audio)

def process_audio(message):
    chat_id = message.chat.id
    subject = user_state[chat_id]["subject"]

    bot.send_message(chat_id, "❗️ Это может занять несколько минут !")

    # Получаем аудио
    if message.voice:
        file_info = bot.get_file(message.voice.file_id)
        file_bytes = bot.download_file(file_info.file_path)
        audio = BytesIO(file_bytes)
        audio_format = "ogg"
    elif message.audio:
        file_info = bot.get_file(message.audio.file_id)
        file_bytes = bot.download_file(file_info.file_path)
        audio = BytesIO(file_bytes)
        audio_format = "mp3"
    else:
        bot.send_message(chat_id, "❗️ Это не аудио файл 😢 Попробуйте снова.")
        bot.register_next_step_handler_by_chat_id(chat_id, process_audio)
        return

    # Конвертация в WAV
    try:
        audio_segment = AudioSegment.from_file(audio, format=audio_format)
        wav_io = BytesIO()
        audio_segment.export(wav_io, format="wav")
        wav_io.seek(0)
    except Exception as e:
        bot.send_message(chat_id, f"⚠️ Ошибка конвертации аудио: {e}")
        return

    # Распознавание речи
    recognizer = sr.Recognizer()
    with sr.AudioFile(wav_io) as source:
        audio_data = recognizer.record(source)
        try:
            text = recognizer.recognize_google(audio_data, language="ru-RU")
        except sr.UnknownValueError:
            bot.send_message(chat_id, "⚠️ Не удалось распознать речь 😢 Попробуйте снова.")
            bot.register_next_step_handler_by_chat_id(chat_id, process_audio)
            return
        except sr.RequestError as e:
            bot.send_message(chat_id, f"⚠️ Ошибка сервиса распознавания: {e}")
            return

    # TXT
    txt_file = BytesIO()
    txt_file.write(text.encode("utf-8"))
    txt_file.seek(0)

    # DOCX
    doc = Document()
    doc.add_paragraph(text)
    doc_file = BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)

    bot.send_document(chat_id, ("{}.txt".format(subject), txt_file))
    bot.send_document(chat_id, ("{}.docx".format(subject), doc_file))
    bot.send_message(chat_id, f"🔊 Аудио преобразовано в текст по 📕 предмету {subject} ✅")

    # После отправки снова показываем расписание
    send_today_schedule(message, user_state[chat_id]["group"])

# ==================== ПОЛЛИНГ ====================
if __name__ == "__main__":
    bot.polling(non_stop=True)
